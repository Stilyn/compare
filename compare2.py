#!/usr/bin/python3

'''
Пример запуска из командной строки
python3 compare2.py Основы.docx Основы2.docx 88
где
    *.docx документы - количество неограничено
    88 - точность совпадения в %
'''

import time
import datetime
import os
# import string
import sys
import docx  # библиотека работа в word
from diff_match_patch import diff_match_patch as diff_module  # для сравнения и раскраски по совету коллег
from docx import Document
# from docx.enum.text import WD_COLOR
from fuzzywuzzy import fuzz
# from jinja2 import Environment, FileSystemLoader
# from nltk.corpus import stopwords
# from nltk.tokenize import word_tokenize
import config
import pandas as pd
import numpy as np
from threading import Thread

# ********************************************   смысловой разбор и поиск ключевых слов
# import pullenti
from pullenti.Sdk import Sdk
from pullenti.ner.ProcessorService import ProcessorService
from pullenti.ner.SourceOfAnalysis import SourceOfAnalysis

# инициализируем pullenti в полном обеме
start_time = time.time()  # время выполнения
Sdk.initialize_all()


def put_to_list(txt, processor, res_list):
    result = processor.process(SourceOfAnalysis(txt))
    for entity in result.entities: res_list.append(str(entity))


# функция генерации ключевых слов
def mind_generate(txt):
    ss = []
    # processor = ProcessorService.create_processor()  # результаты по основным встроенным процессорам pullenti
    # processor_key = ProcessorService.create_specific_processor('KEYWORD')
    # # for analysers in processor_key.analyzers:
    # #    print(analyzers)
    # result = processor_key.process(SourceOfAnalysis(txt))
    # result1 = processor.process(SourceOfAnalysis(txt))
    # # print(result, result1)
    # for entity in result.entities: ss.append(str(entity))  # сделать независимыми потоками
    # for entity1 in result1.entities: ss.append(str(entity1))  # сделать независимыми потоками
    # ss = list(set(ss))  # чистим от дублей

    thr1 = Thread(target=put_to_list, args=(txt, ProcessorService.create_processor(), ss))
    thr2 = Thread(target=put_to_list, args=(txt, ProcessorService.create_specific_processor('KEYWORD'), ss))
    thr = [thr1, thr2];
    for t in thr: t.start(); t.join()
    ss = list(set(ss))  # чистим от дублей
    # print('*** slots **', ss)
    return ss  # возвращает словарь ключевых слов файла


# ********************************************смысловой разбор и поиск ключевых слов

# функция удаления мусорных временных файлов
def spam_del(list_of_files):  # на вход список с названиями файлов для удаления
    for f in list_of_files:
        if os.path.exists(f): os.remove(f)


# функция переименования файлов для формирования временных
def file_rename(file_name):
    n = file_name.split('.')[0] + '_vs' + '.docx'
    return str(n)


# функция удаление пустых параграфов из документа
def strip_file(file, file_new):
    for paragraphs in file.paragraphs:
        if len(paragraphs.text) == 0:
            # delete_paragraph(paragraphs)
            p = paragraphs._element
            p.getparent().remove(p)
            p._p = p._element = None
    file.save(file_new)


# функция добавления пустых абзацев в документ
# def add_par(document, par_count, new_name):
#     for f in range(par_count):
#         document.add_paragraph(' ')
#         document.save(new_name)  # подумать как назвать файл

# функция разбивки русского текста на слова
# def tokenize_ru(file_text):
#     # firstly let's apply nltk tokenization
#     tokens = word_tokenize(file_text)
#
#     # let's delete punctuation symbols
#     tokens = [i for i in tokens if (i not in string.punctuation)]
#
#     # deleting stop_words
#     stop_words = stopwords.words('russian')
#     # stop_words = []
#     stop_words.extend(['что', 'это', 'так', 'вот', 'быть', 'как', 'в', '—', '–', 'к', 'на', '...'])
#     tokens = [i for i in tokens if (i not in stop_words)]
#
#     # cleaning words
#     tokens = [i.replace("«", "").replace("»", "") for i in tokens]
#
#     return tokens

# функция поиска смыслового совпадения параграфов глубина threshold
# def par_match(p1, p2, thresold):
#     dmp = diff_module()
#     diff_module.Match_Threshold = thresold
#     # diff_module.Match_Distance = 0
#     matches = dmp.match_main(p1, p2, 0)
#     return matches

# функция сравнения блоков текста paragraph
# def f_compare(p1, p2):
#     # чистим абзацы от шлака
#     for k in config.symbols_clear:
#         p1 = str.replace(p1, k, ' ')
#         p2 = str.replace(p2, k, ' ')
#     dmp = diff_module()
#     diffs = dmp.diff_main(p1, p2)  # разница
#     # dmp.diff_cleanupSemantic(diffs)
#     # вот здесь нужно значительно улучшить алгоритм сравнения
#     return dmp.diff_prettyHtml(diffs)

# def color_paragraph(paragraph):
#     # paragraphs.runs[s].font.color.rgb = RGBColor(0xff, 0x00, 0x00) # красный текст после нуля просто цвет html
#     # paragraphs.runs[s].font.bold = True # жирный шрифт
#     paragraph.style.font.highlight_color = WD_COLOR.YELLOW  # цвет выделения желтый

# функция выравнивания длин списков
def length_align(list_of_lists):
    ln = []
    for k in list_of_lists:
        ln.append(len(k))
    l = max(ln)
    # print('***' + str(l))
    for m in list_of_lists:
        while len(m) < l:
            m.append(' ')
    return list_of_lists


def par_compare(q1, q2, q4, q5, thresold):
    # q1 q2  - тексты параграфов
    # q4 q5  - ключевые слова параграфов
    q21 = []  # для вывода совпадающих значений 1 и 2 документа
    q2_2 = []  # для вывода несовпадающих значений из 2 документа
    q3 = []  # процент совпадения
    q51 = []
    q5_2 = []  # для вывода несовпадающих значений ключей из 2 документа
    result = []
    for i in range(len(q1)):  # берем все параграфы документа 1
        # q1 сразу выводим в таблицу
        q21.append(' ')
        q51.append(' ')
        q3.append(' ')
        for j in range(len(q2)):
            q2_2.append(' ')
            q5_2.append(' ')
            a = fuzz.WRatio(q1[i], q2[j])  # ищем совпадение по смыслу %
            # a = fuzz.partial_token_sort_ratio(q1[i], q2[j])  # ищем совпадение по словам %
            # print('% текст ********', a)
            b = fuzz.token_sort_ratio(q4[i], q5[j])
            # print('% ключи ********', b)
            q3[i] = str(a) + '|' + str(b)  # или написать процент совпадения как среднее арифметическое из a b
            if int(a) >= int(thresold) and int(b) >= int(thresold) and len(q4[i]) > 0 and len(q5[j]) > 0:
                q21[i] = q2[j]
                q51[i] = q5[j]
            else:  # наполняем мешок с несовпадениями
                q2_2[j] = q2[j]
                q5_2[j] = q5[j]
    # объединяем совпавшие и несовпавшие
    q21.extend(q2_2)
    q51.extend(q5_2)

    # очистить мешок с несовпадениями от пустых значений
    while len(q21) > (len(q1) + len(q2)): del q21[-1]
    while len(q51) > (len(q1) + len(q2)): del q51[-1]

    # выравниваем размерность перед формированием результата
    # ln = max(len(q1), len(q4), len(q3), len(q21), len(q51))
    mass = [q1, q4, q3, q21, q51]
    result = length_align(mass)  # выравниваем длину списков
    return result


DF = []  # список датафреймов файлов документов


# функция разбиения документа формирования ключевых слов и прочей ... для каждого параграфа
def split_doc(file_name, paragraphs):  # , doc_dict)
    indexes = []
    index: int = 0  # добавляем индекс абзаца
    list_text = []  # список текстов параграфов
    list_keywords = []  # список ключевых слов для текстов параграфов
    level = []
    # list_level = []  # список уровней документа
    list_doc_parts = []  # части документа
    for g in paragraphs:  # заранее готовим списки ключевых слов и  тектсов параграфов для документа 1
        indexes.append(index)
        g_mind = mind_generate(g.text)  # keywords
        # print(g.text[0])

        # выясняем level для параграфа
        level_ = ''
        for k in config.doc_levels.items():  # пара (ключ, значение) в словаре
            # print(k[1])  # это список
            for lab in k[1]:
                if lab in str(g.text[:5]):  # подумать пронормальное сравнение
                    level_ = k[0]
        level.append(level_)
        list_text.append(g.text)
        # list_keywords.append(' '.join(g_mind))
        list_keywords.append(list(set(g_mind)))
        list_doc_parts_ = []
        # print(mind_generate(config.doc_parts_list))
        for p in mind_generate(config.doc_parts_list):  # добавляем части документа в датафрейм
            if p in g_mind:
                list_doc_parts_.append(p)
            else:
                list_doc_parts_.append('')
        list_doc_parts.append(' '.join(list_doc_parts_).lstrip(' ').rstrip(' '))  # убираем пробелы в начале и конце
        index += 1
    # готовим датафрейм документа чтобы потом сравнивать

    df = pd.DataFrame(
        {'file': file_name, 'par_index': indexes, 'parent_index': '', 'doc_parts': list_doc_parts, 'level': level, 'text': list_text,
         'keywords': list_keywords})
    DF.append(df)
    # print(df)
    # return df


'''для использования отладочного и боевого режимов'''
if len(sys.argv) > 1:  # если из под командной строки запускаем
    files = []  # перечень файлов для сравнения
    for i in range(1, len(sys.argv) - 1):
        # print('сравниваю ' + sys.argv[1] + ' и ' + sys.argv[2])
        files.append(sys.argv[i])
    thresold = sys.argv[len(sys.argv) - 1]  # сделать вычислением последнего системного аргумента
    print(files, thresold)  # получаем список всех файлов для сравнения и глубину в %

else:
    print('отладочный режим')  # если не из под командной строки запускаем
    # files = ['906_2013.docx', '51_2014.docx', '64_2020.docx']
    files = ['64_2020.docx']
    thresold = config.thresold

# print(config.doc_levels.values()) # проверяем словарь на правильность

print('***** Готовлю ключевые слова *******')
start_time_keys = time.time()  # время начала выполнения

files_vs = []  # список имен файлов сравнения
for i in range(len(files)):
    file_vs = file_rename(files[i])  # делаем временные файлы для сравнения
    # files[i] = docx.Document(files[i])  # линкуем файл  docx
    strip_file(docx.Document(files[i]), file_vs)  # убираем из файлов лишние строки и сохраняем под другими именами
    files_vs.append(file_vs)
    print(len(docx.Document(files[i]).paragraphs))

th = []  # список потоков для многопоточности
for i in files_vs:
    th_ = Thread(target=split_doc, args=(
        i, docx.Document(i).paragraphs))  # формируем ключевые слова для каждого параграфа почищенных файлов
    th.append(th_)
for t in th: t.start(); t.join()  # запучкаем многопотоково формирование датафреймов для  каждого файла
# print(globals()['DF']) # это все датафреймы всех файлов

# объединяем все датафреймы в один
# df = pd.concat(globals()['DF'], keys=files_vs)
df = pd.concat(globals()['DF'])
print("Время выполнения--- %s seconds ---" % (time.time() - start_time_keys) + '\n\n')

print('***** Сравниваю по смыслу, ключевым словам и готовлю сводную таблицу xlsx *******')
start_time_compare = time.time()  # время начала выполнения
file_compare_name_d = str(datetime.datetime.now()).replace(' ', '_').replace(':', '_').split('.')[0] + '.xlsx'
file_compare_name_ht = str(datetime.datetime.now()).replace(' ', '_').replace(':', '_').split('.')[0] + '.html'

'''теперь сравниваем отдельно тексты отдельно ключевые слова делаем pd.series и вставляем в новый датафрейм
ткусты сравниваем по уровням'''
print(df.shape)
print(df.columns)
print(df.index)

# готовим список фильтров уровня
filter_level = []
#for lvl in config.doc_levels.keys(): filter_level.append('df[\'level\'] == ' + lvl)
for lvl in config.doc_levels.keys(): filter_level.append(lvl)
print(filter_level) # список фильтров уровней

# готовим список докпартсов через mind_generate(config.doc_parts_list)
filter_doc_parts = []
for dprt in mind_generate(config.doc_parts_list): filter_doc_parts.append(dprt)
print(filter_doc_parts) # список фильтров докпартс

# делаем отдельный служебный датафрейм
temp_df = []
for i in filter_level:
    flt_lev = df['level'] == i
    for j in filter_doc_parts:
        flt_dp = df['doc_parts'] == j
        temp_df_ = df[['file', 'par_index', 'parent_index', 'doc_parts', 'level']].loc[flt_lev & flt_dp]  # new dataframe contains selected elements
        temp_df.append(temp_df_)
new_df = pd.concat(temp_df) # делаем новый датафрейм с целлями задачами и мероприятиями разбитыми по уровням
#new_df = pd.concat(temp_df, keys=filter_level) # делаем новый датафрейм с целлями задачами и мероприятиями разбитыми по уровням
# не забыть удалить старые датафреймы при необходимости
print(new_df)


# for j in filter_doc_parts:

# df1 = df[['file', 'par_index', 'parent_index', 'doc_parts', 'level']].loc[filter_level & filter_doc_parts]  # new dataframe contains selected elements
# ind = df1['par_index']
# #print(ind)
# i_list = []
# for a in ind: i_list.append(a); # childs par_indexes
# print(i_list)

# print(df[['par_index', 'parent_index', 'level']])

# добавляем parent_index всем дочерним
# for i in range(len(a)):
#     df.at[a[i], 'parent_index'] = a[0]-1  # неправильно ищу значение - придумать как
# print(df[['par_index', 'parent_index', 'level']].loc[a])

# print(df1.shape, df1.columns)
# print(df1)
# print(df[['doc_parts', 'level', 'text']].loc[filter_doc_parts])

# comp = []
# # добавление результатов сравнения
# for w in range(len(files_vs)):
#     comp_ = par_compare(texts[0], texts[w], keywords[0], keywords[w], thresold)
#     comp.append(comp_)  # получаем дложенный список
# length_align(comp)  # выравниваем размеры списков внутри вложенного
# # print(len(comp[0][0]),len(comp[1][0]))
#
# df = pd.DataFrame()
# df[files_vs[0]] = np.array(comp[0][3])
# df[str('keywords' + files_vs[0])] = np.array(comp[0][4])
# for r in range(1, len(comp)):
#     df[str('%' + str(r))] = pd.Series(comp[r][2])
#     df[files_vs[r]] = pd.Series(comp[r][3])
#     df[str('keywords' + files_vs[r])] = pd.Series(comp[r][4])
# print(df)
print("Время выполнения--- %s seconds ---" % (time.time() - start_time_compare) + '\n\n')

''' этот блок потом взять и использовать для записи результтирующих файлов'''
start_time_xlsx = time.time()
print('*****Записываю xlsx*******')
df.to_excel(config.results_folder + file_compare_name_d)  # xlsx
print("Время выполнения--- %s seconds ---" % (time.time() - start_time_xlsx) + '\n\n')

start_time_html = time.time()
print('*****Записываю html*******')
html_string = '''
<html>
  <head>
    <meta charset="utf-8">
    <title>HTML Pandas Dataframe with CSS</title>
</head>
  <link rel="stylesheet" type="text/css" href="df_style.css"/>
  <body>
    {table}
  </body>
</html>
'''
with open(config.results_folder + file_compare_name_ht, 'w') as fh:
    fh.write(html_string.format(table=df.to_html()))
fh.close()

print("Время выполнения--- %s seconds ---" % (time.time() - start_time_html) + '\n\n')

start_time_dlt = time.time()
print('*****Удаляю временные файлы*******')
spam_del(files_vs)
print("Время выполнения--- %s seconds ---" % (time.time() - start_time_dlt) + '\n\n')


print("Общее время выполнения--- %s seconds ---" % (time.time() - start_time))
