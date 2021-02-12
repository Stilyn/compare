#!/usr/bin/python3

# This is a sample Python script.

# Press Shift+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.

import config
import docx  # библиотека работа в word
from docx import Document
from docx.shared import RGBColor
import nltk  # библиотека разбора текста
import string


# print(len(doc1.paragraphs))  # количество абзацев в документе
# print(doc1.paragraphs[0].text)  # текст первого абзаца в документе
# print(doc1.paragraphs[1].text)  # текст второго абзаца в документе
# print(doc1.paragraphs[2].runs[0].text) # текст первого Run второго абзаца

# просто так функция привет )))
# def print_hi(name):
# Use a breakpoint in the code line below to debug your script.
# print(f'Hi, {name}')  # Press Ctrl+F8 to toggle the breakpoint.

# функция переименования файлов для формирования временных
def file_rename(file_name):
    n = file_name.split('.')[0] + '_vs' + '.docx'
    return str(n)


# функция удаление параграфа из документа
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


# функция удаление пустых строк из документа
def strip_file(file, file_new):
    for paragraphs in file.paragraphs:
        if len(paragraphs.text) == 0:
            delete_paragraph(paragraphs)
    file.save(file_new)


# функция добавления пустых абзацев в документ
def add_par(document, par_count, new_name):
    for f in range(par_count):
        document.add_paragraph(' ')
        document.save(new_name)  # подумать как назвать файл


# функция сравнения блоков текста paragraph
def f_compare(p1, p2):
    # чистим абзацы от шлака
    for k in config.symbols_clear:
        p1 = str.replace(p1, k, ' ')
        p2 = str.replace(p2, k, ' ')

    # преобразовать каждый текст в список предложений
    # print(len(p1), len(p2))
    sentences1 = nltk.sent_tokenize(p1)  # массив предложений 1 убираем точнки из них
    sentences2 = nltk.sent_tokenize(p2)  # массив предложений 2 убираем точки из них
    res = list(set(sentences1) ^ set(sentences2))  # результат сравнения - список предложений - ^ - симметричная разность - чего нет хотя бы в одном документе
    return res  # вставить return чтобы работать с результатами как с переменной


# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    # print_hi('PyCharm')  # тестирование работы функций

    doc1 = docx.Document(config.file1)  # линкуем первый файл как docx из конфигурационника
    doc2 = docx.Document(config.file2)  # линкуем второй файл как docx из конфигурационника

    # убрать пустые строки из файлов - это будут выровненные файлы для сравнения
    strip_file(doc1, file_rename(config.file1))  # убираем из файлов лишние строки и сохраняем под другими именами
    strip_file(doc2, file_rename(config.file2))  # убираем из файлов лишние строки и сохраняем под другими именами

    '''
    загружаем очищенные документы
    file_rename(config.file1) это имя переименованного файла 1
    file_rename(config.file2) это имя переименованного файла 2
    '''
    doc1 = docx.Document(file_rename(config.file1))
    doc2 = docx.Document(file_rename(config.file2))

    '''
    выбрать какой из подрезанных длиннее чтобы потом всегда из большего вычитать меньшее
    если больше или равно, то первый иначе второй
    '''
    if len(doc1.paragraphs) >= len(doc2.paragraphs):
        # уравнять количество параграфов путем добавления пустых параграфов в нужный жокумент
        add_par(doc2, (len(doc1.paragraphs) - len(doc2.paragraphs)), file_rename(config.file2))
    else:
        # уравнять количество параграфов путем добавления пустых параграфов в нужный жокумент
        add_par(doc1, (len(doc2.paragraphs) - len(doc1.paragraphs)), file_rename(config.file1))
    ln = len(doc1.paragraphs)
    # print(ln)  # количество абзацев в самом длинном документе

    for i in range(ln):
        '''
        сравниваем тудасюда если вдруг будет вычитание из меньшего массива больший
        и результат будет пустой тогдв надо в обратку чтоб от большего меньший
        поэтому в функции используем симметничную разность ^
        '''
        diff = f_compare(doc1.paragraphs[i].text, doc2.paragraphs[i].text)  # сравниваем по параграфам с добавлением признака документа
        # теперь найти в каком файле эта фраза и подсветить ее
        if len(diff) > 0:
            #print(len(diff), diff)
            print(diff)
            # ищем соответствие в параграфе в обоих документах и раскрашиваем тот где найдем
            # для этого надо новую функцию написать
            for g in range(len(diff)):
                if diff[g].strip() in doc1.paragraphs[i].text.strip():  # strip для удаления пробелов в начале и коце строки
                    for s in range(len(doc1.paragraphs[i].runs)):
                        doc1.paragraphs[i].runs[s].font.color.rgb = RGBColor(0xff, 0x00, 0x00) # после нуля просто цвет html
                        doc1.save(file_rename(config.file1))
                if diff[g].strip() in doc2.paragraphs[i].text.strip():  # strip для удаления пробелов в начале и коце строки
                    for s in range(len(doc2.paragraphs[i].runs)):
                        doc2.paragraphs[i].runs[s].font.color.rgb = RGBColor(0xff, 0x00, 0x00) # после нуля просто цвет html
                        doc2.save(file_rename(config.file2))



