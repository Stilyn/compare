#!/usr/bin/python3

'''
Пример запуска из командной строки
python3 compare.py Основы.docx Основы2.docx 88
'''

import time
import datetime
# import os
# import string
import sys
import docx  # библиотека работа в word
from diff_match_patch import diff_match_patch as diff_module  # для сравнения и раскраски по совету коллег
from docx import Document
# from docx.enum.text import WD_COLOR
from fuzzywuzzy import fuzz
# from jinja2 import Environment, FileSystemLoader
# from nltk.corpus import stopwords
# from nltk.tokenize import word_tokenize
import config
import pandas as pd
import numpy as np
from threading import Thread
from multiprocessing import Process

# ********************************************   смысловой разбор и поиск ключевых слов
# import pullenti
from pullenti.Sdk import Sdk
from pullenti.ner.ProcessorService import ProcessorService
from pullenti.ner.SourceOfAnalysis import SourceOfAnalysis

# from pullenti.ner.AnalysisResult import AnalysisResult
# from pullenti.ner.Analyzer import Analyzer
# from pullenti.ner.ExtOntology import ExtOntology
# from pullenti.ner.ExtOntologyItem import ExtOntologyItem
# from pullenti.ner.MetaToken import MetaToken
# from pullenti.ner.MorphCollection import MorphCollection
# from pullenti.ner.NumberSpellingType import NumberSpellingType
# from pullenti.ner.NumberToken import NumberToken
# from pullenti.ner.Processor import Processor
# from pullenti.ner.ProxyReferent import ProxyReferent
# from pullenti.ner.Referent import Referent
# from pullenti.ner.Slot import Slot
# from pullenti.ner.TextAnnotation import TextAnnotation
# from pullenti.ner.TextToken import TextToken
# from pullenti.ner.Token import Token
# from pullenti.ner.keyword import KeywordAnalyzer
# инициализируем pullenti в полном обеме
start_time = time.time()  # время выполнения
Sdk.initialize_all()


# функция генерации ключевых слов
def mind_generate(txt):
    ss = []
    processor = ProcessorService.create_processor()  # результаты по основным встроенным процессорам pullenti
    processor_key = ProcessorService.create_specific_processor('KEYWORD')
    # for analysers in processor_key.analyzers:
    #    print(analyzers)
    result = processor_key.process(SourceOfAnalysis(txt))
    result1 = processor.process(SourceOfAnalysis(txt))
    # print(result, result1)
    for match in result.entities: ss.append(str(match))  # сделать независимыми потоками
    for match1 in result1.entities: ss.append(str(match1))  # сделать независимыми потоками
    ss = list(set(ss))  # чистим от дублей
    # print('*** slots **', ss)
    return ss  # возвращает словарь ключевых слов файла

# ********************************************смысловой разбор и поиск ключевых слов

# функция переименования файлов для формирования временных
def file_rename(file_name):
    n = file_name.split('.')[0] + '_vs' + '.docx'
    return str(n)


# функция удаление пустых параграфов из документа
def strip_file(file, file_new):
    for paragraphs in file.paragraphs:
        if len(paragraphs.text) == 0:
            # delete_paragraph(paragraphs)
            p = paragraphs._element
            p.getparent().remove(p)
            p._p = p._element = None
    file.save(file_new)


# функция добавления пустых абзацев в документ
# def add_par(document, par_count, new_name):
#     for f in range(par_count):
#         document.add_paragraph(' ')
#         document.save(new_name)  # подумать как назвать файл

# функция разбивки русского текста на слова
# def tokenize_ru(file_text):
#     # firstly let's apply nltk tokenization
#     tokens = word_tokenize(file_text)
#
#     # let's delete punctuation symbols
#     tokens = [i for i in tokens if (i not in string.punctuation)]
#
#     # deleting stop_words
#     stop_words = stopwords.words('russian')
#     # stop_words = []
#     stop_words.extend(['что', 'это', 'так', 'вот', 'быть', 'как', 'в', '—', '–', 'к', 'на', '...'])
#     tokens = [i for i in tokens if (i not in stop_words)]
#
#     # cleaning words
#     tokens = [i.replace("«", "").replace("»", "") for i in tokens]
#
#     return tokens

# функция поиска смыслового совпадения параграфов глубина threshold
# def par_match(p1, p2, thresold):
#     dmp = diff_module()
#     diff_module.Match_Threshold = thresold
#     # diff_module.Match_Distance = 0
#     matches = dmp.match_main(p1, p2, 0)
#     return matches

# функция сравнения блоков текста paragraph
# def f_compare(p1, p2):
#     # чистим абзацы от шлака
#     for k in config.symbols_clear:
#         p1 = str.replace(p1, k, ' ')
#         p2 = str.replace(p2, k, ' ')
#     dmp = diff_module()
#     diffs = dmp.diff_main(p1, p2)  # разница
#     # dmp.diff_cleanupSemantic(diffs)
#     # вот здесь нужно значительно улучшить алгоритм сравнения
#     return dmp.diff_prettyHtml(diffs)

# def color_paragraph(paragraph):
#     # paragraphs.runs[s].font.color.rgb = RGBColor(0xff, 0x00, 0x00) # красный текст после нуля просто цвет html
#     # paragraphs.runs[s].font.bold = True # жирный шрифт
#     paragraph.style.font.highlight_color = WD_COLOR.YELLOW  # цвет выделения желтый

# функция формирования датасета сравнения параграфов

def length_align(list_of_lists):
    ln = []
    for k in list_of_lists:
        ln.append(len(k))
    l = max(ln)
    # print('***' + str(l))
    for m in list_of_lists:
        while len(m) < l:
            m.append(' ')
    return list_of_lists


def par_compare(q1, q2, q4, q5, thresold):
    # q1 q2  - тексты параграфов
    # q4 q5  - ключевые слова параграфов
    q21 = []  # для вывода совпадающих значений 1 и 2 документа
    q2_2 = []  # для вывода несовпадающих значений из 2 документа
    q3 = []  # процент совпадения
    q51 = []
    q5_2 = []  # для вывода несовпадающих значений ключей из 2 документа
    result = []
    for i in range(len(q1)):  # берем все параграфы документа 1
        # q1 сразу выводим в таблицу
        q21.append(' ')
        q51.append(' ')
        q3.append(' ')
        for j in range(len(q2)):
            q2_2.append(' ')
            q5_2.append(' ')
            # необходимо разбить на потоки
            # th1 = Thread(target=fuzz.WRatio, args=(q1[i], q2[j]))
            # th2 = Thread(target=fuzz.token_sort_ratio, args=(q4[i], q5[j]))
            # th1.start()
            # th2.start()
            # th1.join()
            # th2.join()
            a = fuzz.WRatio(q1[i], q2[j])  # ищем совпадение по смыслу %
            # a = fuzz.partial_token_sort_ratio(q1[i], q2[j])  # ищем совпадение по словам %
            # print('% текст ********', a)
            b = fuzz.token_sort_ratio(q4[i], q5[j])
            # print('% ключи ********', b)
            q3[i] = str(a) + '|' + str(b)  # или написать процент совпадения как среднее арифметическое из a b
            if int(a) >= int(thresold) and int(b) >= int(thresold) and len(q4[i]) > 0 and len(q5[j]) > 0:
                q21[i] = q2[j]
                q51[i] = q5[j]
            else:  # наполняем мешок с несовпадениями
                q2_2[j] = q2[j]
                q5_2[j] = q5[j]
    # объединяем совпавшие и несовпавшие
    q21.extend(q2_2)
    q51.extend(q5_2)

    # очистить мешок с несовпадениями от пустых значений
    while len(q21) > (len(q1) + len(q2)): del q21[-1]
    while len(q51) > (len(q1) + len(q2)): del q51[-1]

    # выравниваем размерность перед формированием результата
    # ln = max(len(q1), len(q4), len(q3), len(q21), len(q51))
    mass = [q1, q4, q3, q21, q51]
    result = length_align(mass)  # выравниваем длину списков
    return result


# функция разбиения документа формирования ключевых слов для каждого параграфа
def split_doc(paragraphs, list_text, list_keywords):
    for g in paragraphs:  # заранее готовим списки ключевых слов и  тектсов параграфов для документа 1
        g_mind = mind_generate(g.text)
        list_text.append(g.text)
        list_keywords.append(' '.join(g_mind))

        # вот здесь возможно переписать на словари через dict.fromkeys
        # и потом из этого сформировать датафрейм сразу
        # a = [1, 2, 3, ]
        # b = ['ss', 'ddd', 'sdadas']
        # s = dict.fromkeys(a, b)
        # print(s)
        # {1: ['ss', 'ddd', 'sdadas'], 2: ['ss', 'ddd', 'sdadas'], 3: ['ss', 'ddd', 'sdadas']}
        # doc_dict = dict.fromkeys(list_keywords,list_text)
        # print(doc_dict)
        #return doc_dict


'''для использования отладочного и боевого режимов'''
if len(sys.argv) > 1:  # если из под командной строки запускаем
    files = []  # перечень файлов для сравнения
    for i in range(1, len(sys.argv) - 1):
        # print('сравниваю ' + sys.argv[1] + ' и ' + sys.argv[2])
        files.append(sys.argv[i])
    thresold = sys.argv[len(sys.argv) - 1]  # сделать вычислением последнего системного аргумента
    print(files, thresold)  # получаем список всех файлов для сравнения и глубину в %

else:
    print('отладочный режим')  # если не из под командной строки запускаем
    files = ['906_2013.docx', '51_2014.docx', '64_2020.docx']
    thresold = config.thresold


print('***** Готовлю ключевые слова *******')
start_time_keys = time.time()  # время начала выполнения

texts = []  # абзацы документа
keywords = []  # ключевые слова абзацев документа
files_vs = []  # список имен файлов сравнения
for i in range(len(files)):
    file_vs = file_rename(files[i])  # делаем временные файлы для сравнения
    files[i] = docx.Document(files[i])  # линкуем файл  docx
    strip_file(files[i], file_vs)  # убираем из файлов лишние строки и сохраняем под другими именами
    files_vs.append(file_vs)
    print(len(files[i].paragraphs))
    texts_ = []
    keywords_ = []
    # pc1 = Process(target=split_doc, args=(files[i].paragraphs, texts_, keywords_))
    # pc1.start()
    # pc1.join()
    #
    split_doc(docx.Document(files_vs[i]).paragraphs, texts_, keywords_)  # формируем ключевые слова для каждого
    # параграфа почищенных файлов

    # сделать через на multiprocessing
    # th1 = Thread(target=split_doc, args=(doc1.paragraphs, q1, q4))  # поток 1
    # th2 = Thread(target=split_doc, args=(doc2.paragraphs, q2, q5))  # поток 2
    # th1.start()
    # th2.start()
    # th1.join()
    # th2.join()
    texts.append(texts_)
    keywords.append(keywords_)

print("Время выполнения--- %s seconds ---" % (time.time() - start_time_keys) + '\n\n')

print('***** Сравниваю по смыслу, ключевым словам и готовлю сводную таблицу xlsx *******')
start_time_compare = time.time()  # время начала выполнения
file_compare_name_d = str(datetime.datetime.now()).replace(' ', '_').replace(':', '_').split('.')[0] + '.xlsx'
file_compare_name_ht = str(datetime.datetime.now()).replace(' ', '_').replace(':', '_').split('.')[0] + '.html'

comp = []
# добавление результатов сравнения
for w in range(len(files_vs)):
    comp_ = par_compare(texts[0], texts[w], keywords[0], keywords[w], thresold)
    comp.append(comp_)  # получаем дложенный список
length_align(comp)  # выравниваем размеры списков внутри вложенного
# print(len(comp[0][0]),len(comp[1][0]))

df = pd.DataFrame()
df[files_vs[0]] = np.array(comp[0][3])
df[str('keywords' + files_vs[0])] = np.array(comp[0][4])
for r in range(1, len(comp)):
    df[str('%' + str(r))] = pd.Series(comp[r][2])
    df[files_vs[r]] = pd.Series(comp[r][3])
    df[str('keywords' + files_vs[r])] = pd.Series(comp[r][4])
print(df)
print("Время выполнения--- %s seconds ---" % (time.time() - start_time_compare) + '\n\n')

start_time_xlsx = time.time()
print('*****Записываю xlsx*******')
df.to_excel(file_compare_name_d)  # xlsx
print("Время выполнения--- %s seconds ---" % (time.time() - start_time_xlsx) + '\n\n')

start_time_html = time.time()
print('*****Записываю html*******')
# df.to_html(file_compare_name_ht, encoding='utf-8')  # html table
html_string = '''
<html>
  <head>
    <meta charset="utf-8">
    <title>HTML Pandas Dataframe with CSS</title>
</head>
  <link rel="stylesheet" type="text/css" href="df_style.css"/>
  <body>
    {table}
  </body>
</html>
'''
with open(file_compare_name_ht, 'w') as fh:
    fh.write(html_string.format(table=df.to_html()))
fh.close()

print("Время выполнения--- %s seconds ---" % (time.time() - start_time_html) + '\n\n')

print("Общее время выполнения--- %s seconds ---" % (time.time() - start_time))
